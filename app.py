import streamlit as st
import os
from dotenv import load_dotenv
import google.generativeai as genai
from docx import Document
from io import BytesIO
import json
import latex2mathml.converter # 설치하신 라이브러리 import

# ==========================================
# [설정] API 키를 입력하세요
# ==========================================
load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=api_key) 

# ==========================================
# [함수 1] 문제 생성 (Gemini 2.5 Flash)
# ==========================================
def generate_problems(topic, difficulty, count):
    # 최신 모델 사용
    model = genai.GenerativeModel('gemini-2.5-flash')
    
    prompt = f"""
    당신은 한국의 고등수학 문제 출제 위원입니다.
    다음 조건에 맞춰 수학 문제를 만들고, 반드시 Python 리스트 형식의 JSON 데이터만 반환하세요.
    (마크다운 코드 블록이나 잡담을 절대 포함하지 마세요.)

    [출제 조건]
    1. 과목/단원: 공통수학2 > {topic}
    2. 난이도: {difficulty}
    3. 문제 수: {count}개
    4. 형식: 객관식 또는 단답형
    
    [수식 표현 규칙 - LaTeX]
    - 수식은 표준 LaTeX 문법을 사용하세요. (예: \\frac{{1}}{{2}}, x^2)
    - 중요: JSON 문자열 안에서 백슬래시(\\)가 깨지지 않도록 주의하세요.

    [출력 포맷 예시 (JSON)]
    [
        {{
            "q_num": 1,
            "question_text": "집합 A={{1, 2}}일 때...", 
            "answer": "3",
            "solution": "A의 부분집합 개수는 2^2=4이다."
        }}
    ]
    """
    
    try:
        response = model.generate_content(prompt)
        # JSON 파싱 (실수 방지 로직)
        text = response.text.replace("```json", "").replace("```", "").strip()
        # 혹시 모를 앞뒤 공백이나 잡다한 텍스트 제거 시도
        if "[" in text and "]" in text:
            start = text.find("[")
            end = text.rfind("]") + 1
            text = text[start:end]
            
        return json.loads(text)
    except Exception as e:
        st.error(f"생성 실패 (API 또는 파싱 오류): {e}")
        st.write("--- 원본 응답 ---")
        st.code(response.text) # 디버깅용
        return []

# ==========================================
# [함수 2] Word 파일 생성
# ==========================================
def create_docx(problems):
    doc = Document()
    doc.add_heading('공통수학2 문제은행', 0)
    doc.add_paragraph('Created by AI Math Bank\n')
    
    for p in problems:
        # 1. 문제
        p_para = doc.add_paragraph()
        run_num = p_para.add_run(f"{p['q_num']}. ")
        run_num.bold = True
        run_num.font.size = 140000 # 폰트 크기 조절 예시
        
        # 문제 텍스트 (LaTeX 코드가 그대로 들어감 -> 워드에서 변환 필요)
        p_para.add_run(p['question_text'])
        doc.add_paragraph(" ") # 공백
        
    doc.add_page_break()
    doc.add_heading('정답 및 해설', 1)
    
    for p in problems:
        s_para = doc.add_paragraph()
        s_para.add_run(f"{p['q_num']}. 정답: {p['answer']}\n").bold = True
        s_para.add_run(f"해설: {p['solution']}")
        
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# [UI] Streamlit 화면 구성 (여기가 복구됨!)
# ==========================================
st.set_page_config(page_title="Math Question Bank", layout="wide")

st.title("🧮 AI 수학 문제은행 (Pro)")
st.caption("Gemini 2.5 Flash | LaTeX Support")

# 사이드바 (옵션 선택창)
with st.sidebar:
    st.header("출제 옵션")
    selected_topic = st.selectbox(
        "단원 선택",
        ["집합의 뜻과 표현", "집합의 연산", "명제", "함수", "유리함수와 무리함수"]
    )
    difficulty = st.select_slider(
        "난이도",
        options=["하", "중", "상", "최상"]
    )
    count = st.number_input("문제 수", min_value=1, max_value=10, value=3)
    
    st.divider()
    generate_btn = st.button("문제 생성하기", type="primary")

# 메인 화면 로직
if generate_btn:
    with st.spinner('Gemini가 문제를 출제 중입니다...'):
        problems = generate_problems(selected_topic, difficulty, count)
        
        if problems:
            st.success(f"총 {len(problems)}개의 문제가 생성되었습니다!")
            
            # 탭으로 문제와 해설 분리
            tab1, tab2 = st.tabs(["📄 문제 미리보기", "📝 해설 보기"])
            
            with tab1:
                for p in problems:
                    st.markdown(f"#### 문제 {p['q_num']}")
                    # 웹에서는 LaTeX를 렌더링해서 보여줌 (가독성 UP)
                    st.latex(p['question_text']) 
                    st.divider()
            
            with tab2:
                for p in problems:
                    st.markdown(f"**{p['q_num']}. 정답: {p['answer']}**")
                    st.write(p['solution']) # 해설도 LaTeX가 있다면 st.latex() 사용 가능
                    st.divider()
            
            # 다운로드 버튼
            docx_file = create_docx(problems)
            st.download_button(
                label="📥 Word 파일 다운로드",
                data=docx_file,
                file_name=f"math_{selected_topic}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )